'''
Created on Sep 26, 2022

Last Edited on Nov 16, 2022

@author: acrie
'''

import os
import tkinter as tk
from tkinter import *
from tkinter import filedialog as fd
from tkinter import ttk
from tkinter.filedialog import askdirectory as ad
from tkinter import messagebox
import PIL
from PIL import Image
import docx
from docx import Document
import json


global wW
wW = 950 #window Width
wWstr = str(wW)
root = Tk()
root.geometry(wWstr + "x750") #width 950 height 750
root.resizable(False, False)
root.title("Filosopher's Stone")

bgframe = Frame(root, bg = 'yellow')
bgframe.place(width = wW, height = 750)

my_details = {
    "set1" : "1",
    "fts1": "img",
    "dts1": "folderselect",
    "seldes1": "none"
    }

#set = whole set in a row
#fts = file type selected: image (img) or text (txt) types (switch button)(is for file browser and the end type option menu)
#dts = destination type selection: folder or application (which could itself be the destination or sends the converted files to their intended final location)
#seldes = selected destination ex. picture folder
#delorig = delete original files? 'Yes' deletes original file after created a converted copy. 'No' keeps original files after creating converted copy.
#convert = conversion button

with open('fileconverterdata.json', 'w') as f:
    json.dump(my_details, f)

global ypad
ypad = 7
global xpad
xpad = 15
global total_sets
setnumdict = {} #holds number string for functions
framedict = {} #frame dictionary
delbtndict = {} #section delete button
ftsbtndict = {} #fts = file type switch
ofdbtndict = {} #ofd = open file directory
filelistdict = {} #temporarily stores list of files for conversion's sake (not saved in json)
dftmenudict = {} #dft = destination file type
dftoptdict = {} #stores dft options, includes selected
dtsdict = {} #des = destination type switch
dcbtndict = {} #dc = destination chooser
seldesdict = {} #selected destination
convertdict = {} #convert buttons



img_type_list = [".jpg", ".jpeg",".png", ".gif", ".dds"]

txt_type_list = [".txt",".docx", ".doc"] #future support for excel and pdf files in the works


def updatejson():
    with open('fileconverterdata.json', 'w') as f:
        json.dump(obj, f)

def select_txtfiles(n):
    txtfiles = fd.askopenfilenames(
        title='Select text files',
        initialdir='/',
        filetypes=(
        ('All Supported Text Files', '*.txt .docx .doc'),
        ('Txt Files','*.txt'),
        ('Word Documents','*.docx .doc')
        )
    )
    
    addtolist(n, txtfiles)

def select_imgfiles(n):
    imgfiles = fd.askopenfilenames(
        title='Select image files',
        initialdir='/',
        filetypes=(
        ('All Supported Image Files', '*.jpg .jpeg .png .gif .dds'),
        ('JPEG/JPG Files','*.jpg .jpeg'),
        ('PNG Files','*.png'),
        ('GIF Files','*.gif'),
        ('DDS Files','*.dds')
        )
    )
    
    addtolist(n, imgfiles)


def addtolist(n, f): #f here equals files, not frame
    fl = "set" + n
    filelistdict[fl].clear()
    for x in f:
        filelistdict[fl].append(x)

def select_folder(n):
    dest = "set" + n
    folderdir = ad(parent=root,initialdir="/",title='Please select a directory')
    seldesdict[dest] = folderdir
    obj["seldes" + n] = seldesdict[dest]
    updatejson()


    
#switches options for the file directory and for the desired file type drop down menu from image file types to text file types and vice versa
def imgtxtswitch(n):
    f = "fts" + n
    s = "set" + n
    if obj[f] == "img":
        obj[f] = "txt"
        ofdbtndict[s].configure(command = lambda n=n: select_txtfiles(n))
        alteroptions(n, txt_type_list)
    else:
        obj[f] = "img"
        ofdbtndict[s].configure(command = lambda n=n: select_imgfiles(n))
        alteroptions(n, img_type_list)
    updatejson()
    

def alteroptions(n, l):#l is list
    fft = "set" + n
    dftoptdict[fft].set('')
    dftmenudict[fft]['menu'].delete(0,'end')
    for opt in l: 
        dftmenudict[fft]['menu'].add_command(label=opt, command=tk._setit(dftoptdict[fft], opt))
    dftoptdict[fft].set(l[0]) # default value
    
def desttypeswitcher(n): #currently only set up to allow for folder destination, other destination types in the works
    dest = 'dts' + n
    if obj[dest] == "folderselect":
        obj[dest] = "folderselect"

def makexbutton(n, f):
    standinxpic = PhotoImage(file = r'C:\Users\acrie\OneDrive\Pictures\Camera Roll\xbuttonstandin.png')
    cset = "set" + n
    if obj[cset] == "1":
        delbtndict[cset] = tk.Label(f, image = standinxpic,height= 1, width=5)
    else:
        delbtndict[cset] = tk.Button(f, text = "X",command= lambda n=n: deleteset(n), height= 1, width=5)
        
    delbtndict[cset].grid(row = 0, column = 1, padx = xpad, pady=ypad)

def makeftswitch(n, f):
    cset = "set" + n
    ftsbtndict[cset] = tk.Button(f, text='Image Files', bg = 'blue', command= lambda n=n:imgtxtswitch(n),height= 1, width=10)
    ftsbtndict[cset].grid(row = 0, column = 2, padx = xpad, pady=ypad)

def makefbbtn(n, f):
    cset = "set" + n
    if obj['fts' + n] == "img":
        ofdbtndict[cset] = tk.Button(f, text='Select File(s)', bg= '#FFD700', command= lambda n=n: select_imgfiles(n),height= 1, width=10)
    else:
        ofdbtndict[cset] = tk.Button(f, text='Select File(s)', bg= '#FFD700', command=lambda n=n: select_txtfiles(n),height= 1, width=10)
        
    ofdbtndict[cset].grid(row = 0, column = 3, padx = xpad, pady=ypad)
    filelistdict[cset] = []    

def makefinalfiletypechooser(n, f):
    cset = "set" + n
    dftoptdict[cset] = tk.StringVar(root)
    if obj["fts"+ n] == "img":
        dftoptdict[cset].set(img_type_list[0])
        dftmenudict[cset]= tk.OptionMenu(f, dftoptdict[cset],*img_type_list) #drop down menu for intended destination file type
    else:
        dftoptdict[cset].set(txt_type_list[0])
        dftmenudict[cset]= tk.OptionMenu(f, dftoptdict[cset],*txt_type_list) #drop down menu for intended destination file type
    
    dftmenudict[cset].config(bg="white", activebackground='#00acdb', highlightthickness= 0 ,width=5)
    dftmenudict[cset]['menu'].config(bg="white", activebackground='#00acdb')
    dftmenudict[cset].grid(row = 0, column = 4, padx = xpad, pady=ypad)


    
def makedesttypeswitch(n, f):
    cset = "set" + n
    dtsdict[cset] = tk.Button(f, text='Local', bg = 'blue', command= lambda n=n: desttypeswitcher(n),height= 1, width=10) #destination type switch
    dtsdict[cset].grid(row = 0, column = 5, padx = xpad, pady=ypad)



def makedestchooser(n, f):
    cset = "set" + n
    dcbtndict[cset] = tk.Button(f, text='Select Destination', bg= '#FFD700', command=lambda n=n:select_folder(n),height= 1, width=15) #destination
    dcbtndict[cset].grid(row = 0, column = 6, padx = xpad, pady=ypad)
    seldesdict[cset] = obj["seldes" + n]
    
    
def makeconverter(n, f):
    cb = "set" + n
    convertdict[cb] = tk.Button(f, text='Convert', bg= '#FFD700', command=lambda n=n:convertcheck(n),height= 1, width=10)
    convertdict[cb].grid(row = 0, column = 7, padx = xpad, pady=ypad)
    
    
def convertcheck(n):
    cset = "set" + n
    d = seldesdict[cset].replace("/","\\") #changes / to \
    ext = dftoptdict[cset].get() #intended file type extension
    if (len(filelistdict[cset]) == 0) and (obj["seldes"+n] == "none"):
        messagebox.showerror("Error", "No files or destination have been selected.")
    elif obj["seldes"+n] == "none":
        messagebox.showerror("Error", "No destination has been selected.")
    elif len(filelistdict[cset]) == 0:
        messagebox.showerror("Error", "No files have been selected.")
    else:
        if obj["fts"+n] == "img":
            convertImg(n, d, cset, ext)
        else:
            convertTxt(n, d, cset, ext)
    messagebox.showinfo("showinfo", "Conversion complete")
    

def convertImg(n, d, cset, ext):
    if obj["dts" + n] == "folderselect":
        for file in filelistdict[cset]:
            s1 = file.replace("/","\\") #changes / to \
            s2 = s1.rfind("\\") #finds location of final \
            s3 = s1.rfind(".") #finds locations of .
            s4 = s1[s2:s3] #separates filename from directory and removes the extension and period
            
            final = d + s4 + ext #makes directory string for saving the new file
            
            im = Image.open(r''+s1+'')
            im.save(r''+final+'')

def convertTxt(n, d, cset, ext):
    if obj["dts" + n] == "folderselect":
        for file in filelistdict[cset]:
            s1 = file.replace("/","\\") #changes / to \
            s2 = s1.rfind("\\") #finds location of final \
            s3 = s1.rfind(".") #finds locations of .
            s4 = s1[s2:s3] #separates filename from directory and removes the extension and period
            
            if ext.startswith('doc'):
                if file.endswith('.txt'):
                    txt2docs(s1, s4, ext, d)
                if (s1[(s3+1):]).startswith('doc'):
                    docs2docs(s1, d, s4, ext)
            
            if ext == ".txt" and (s1[(s3+1):]).startswith('doc'):
                docs2txt(s1, d, s4, ext)
    

            

def txt2docs(s1, d, s4, ext):
    doc = Document()
    final = d + s4 + ext     
    with open(r''+ s1 +'') as f:
        lines = f.readlines()

        for x in lines:
            doc.add_paragraph(x.strip())
    
    doc.save(r''+ final +'')

def docs2txt(s1, d, s4, ext):
    doc = Document(r''+ s1 +'')
    final = d + s4 + ext     
    allText = []
    p = ""
    for docpara in doc.paragraphs:
        allText.append(docpara.text)
        p = p + docpara.text + '\n'
        with open(r''+ final +'', 'w') as f:
            f.write(p)

def docs2docs(s1, d, s4, ext): #doc to docx and vice versa
    doc = Document(r''+ s1 +'')
    final = d + s4 + ext
    doc.save(r''+ final +'')

def makeset(num):
    global total_sets
    f = "set" + num
    #setnumdict[f] = num
    currset = int(num)
    yposit = setYpos(currset)
    framedict[f] = Frame(root, bg = '#00acdb')
    framedict[f].place(x= 5, y= yposit, width = (wW-10), height = 40)
    makexbutton(num, framedict[f])
    makeftswitch(num, framedict[f])
    makefbbtn(num, framedict[f])
    makefinalfiletypechooser(num, framedict[f])
    makedesttypeswitch(num, framedict[f])
    makedestchooser(num, framedict[f])
    makeconverter(num, framedict[f])

def setYpos(n):
    ypos = 150 + (48*(n-1))
    return ypos

def iterthruobj(obj):
    global total_sets
    
    total_sets = 0
    for s in obj:
        if s.startswith("set"):
            total_sets = total_sets+1
            num = obj[s]
            makeset(num)

with open('fileconverterdata.json') as f:
    obj = json.load(f)
    
    iterthruobj(obj)
    



def addnewset():
    global total_sets
    total_sets = total_sets + 1
    ntsstr = str(total_sets) #new total set string
    obj["set" + ntsstr] = ntsstr
    obj["fts" + ntsstr]= "img"
    obj["dts" + ntsstr]= "folderselect"
    obj["seldes" + ntsstr]= "none"
    updatejson()
    makeset(ntsstr)

def deleteset(n):
    global total_sets
    
    num = int(n)
    dif = total_sets-num
    i = total_sets
    
    if dif >= 1:
        destroyset(num)
        while i > num:
            m = num-1
            print(i)
            i=i-1
    else:
        destroyset(num)
        dictpop(num)
    print('end')
    total_sets = total_sets -1


def destroyset(n):
    cset = "set" + str(n)
    framedict[cset].destroy()
    delbtndict[cset].destroy()
    ftsbtndict[cset].destroy()
    ofdbtndict[cset].destroy()
    dftmenudict[cset].destroy()
    dtsdict[cset].destroy()
    dcbtndict[cset].destroy()
    convertdict[cset].destroy()
                
def dictpop(n):
    nstr = str(n)
    cset = "set" + nstr
    framedict.pop(cset)
    delbtndict.pop(cset)
    ftsbtndict.pop(cset)
    ofdbtndict.pop(cset)
    filelistdict.pop(cset)
    dftmenudict.pop(cset)
    dftoptdict.pop(cset)
    dtsdict.pop(cset)
    dcbtndict.pop(cset)
    seldesdict.pop(cset)
    convertdict.pop(cset)

#this shifts the positions of the sets after(below) the deleted set to "fill the created empty space" if necessary 
#i.e. if there are 5 sets and set 4 is deleted, set 5's data and visual widget are moved to set 4's previous position on screen and in the dictionaries
def shiftset(n,m):#unfinshed and in progress
    nstr = str(n)
    mstr = str(m)
    cset = "set" + nstr
    ncset = "set" + mstr
    obj["set" + mstr] = obj["set" + nstr]
    obj["fts" + mstr] = obj["fts" + nstr]
    obj["dts" + mstr] = obj["dts" + nstr]
    obj["seldes" + mstr] = obj["seldes" + nstr]
    framedict[ncset] = framedict[cset]
    delbtndict[ncset] = delbtndict[cset]
    ftsbtndict[ncset] = ftsbtndict[cset]
    ofdbtndict[ncset] = ofdbtndict[cset]
    dftmenudict[ncset] = dftmenudict[cset]
    dtsdict[ncset] = dtsdict[cset]
    dcbtndict[ncset] = dcbtndict[cset]
    convertdict[ncset] = convertdict[cset]
    
    if obj["fts"+ mstr] == "img":
        ofdbtndict[ncset].configure(framedict[ncset], command= lambda mstr=mstr: select_imgfiles(mstr))




aboveframe = Frame(root, bg = 'red')
aboveframe.place(x= 5, y= 2, width = (wW-10), height = 140)

addbtn = tk.Button(aboveframe, text = "Add set", command= addnewset, height = 2, width = 15)
addbtn.grid(row = 0, column = 1, padx = xpad, pady=ypad)


root.mainloop()
